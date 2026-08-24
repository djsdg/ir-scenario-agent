from __future__ import annotations

import json
from pathlib import Path


SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".json", ".docx", ".pdf"}


def read_document(path: str | Path) -> str:
    """Read a requirement document into plain text for the IR extraction step.

    Text and JSON are dependency-free.  DOCX and PDF support is optional so
    the core agent remains small; a clear installation hint is returned when
    one of those adapters is requested without its extra dependency.
    """

    document_path = Path(path)
    suffix = document_path.suffix.casefold()
    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ValueError(f"不支持的文档类型 {suffix or '<无扩展名>'}，支持：{supported}")
    if suffix in {".txt", ".md", ".markdown", ".rst"}:
        return document_path.read_text(encoding="utf-8")
    if suffix == ".json":
        payload = json.loads(document_path.read_text(encoding="utf-8"))
        return json.dumps(payload, ensure_ascii=False, indent=2)
    if suffix == ".docx":
        return _read_docx(document_path)
    return _read_pdf(document_path)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("读取 DOCX 需要可选依赖，请执行：pip install -e \".[docs]\"") from exc

    document = Document(path)
    chunks: list[str] = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("读取 PDF 需要可选依赖，请执行：pip install -e \".[docs]\"") from exc

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)
